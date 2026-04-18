# ------------------------------------------------------------------------------
# SPDX-License-Identifier: Proprietary
# ------------------------------------------------------------------------------

"""Portable DOCX report generation for Silica-X payloads."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from core.artifacts.charts import build_chart_images


def _safe_rows(value: object) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    return [row for row in value if isinstance(row, dict)]


def _set_cell_text(cell: Any, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold


def _add_table(document: Any, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        _set_cell_text(header_cells[index], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            _set_cell_text(cells[index], value)


def _summary_rows(payload: dict[str, Any]) -> list[list[str]]:
    summary = payload.get("summary", {}) if isinstance(payload.get("summary"), dict) else {}
    return [[str(key), str(value)] for key, value in summary.items()]


def generate_docx_report(path: Path, payload: dict[str, Any]) -> str:
    """Write a rich DOCX case report for a full payload."""

    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("python-docx is required for DOCX report output.") from exc

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = f"Silica-X Reporter - {payload.get('target', 'target')}"

    title = document.add_heading("Silica-X Reporter", 0)
    title.runs[0].font.size = Pt(22)
    document.add_paragraph(
        f"Target: {payload.get('target', '-')}\n"
        f"Mode: {(payload.get('metadata') or {}).get('mode', '-')}\n"
        f"Generated: {(payload.get('metadata') or {}).get('generated_at_utc', '-')}"
    )

    document.add_heading("Reporter Brief", level=1)
    document.add_paragraph(str(payload.get("narrative") or "No Reporter Brief was generated."))

    summary_rows = _summary_rows(payload)
    if summary_rows:
        document.add_heading("Overview", level=1)
        _add_table(document, ["Metric", "Value"], summary_rows)

    chart_dir, charts = build_chart_images(payload)
    try:
        if charts:
            document.add_heading("Visual Intelligence", level=1)
            for title_text, key in (
                ("Result Status Distribution", "status_bar"),
                ("Issue Severity Mix", "severity_pie"),
                ("Response Time Trend", "response_line"),
                ("Confidence Histogram", "confidence_hist"),
            ):
                chart_path = charts.get(key)
                if chart_path is None:
                    continue
                document.add_paragraph(title_text)
                document.add_picture(str(chart_path), width=Inches(6.2))

        selected_plugins = payload.get("selected_plugins", []) if isinstance(payload.get("selected_plugins"), list) else []
        selected_filters = payload.get("selected_filters", []) if isinstance(payload.get("selected_filters"), list) else []
        attached_modules = _safe_rows(payload.get("attached_modules"))
        if selected_plugins or selected_filters or attached_modules:
            document.add_heading("Attachables", level=1)
            if selected_plugins:
                document.add_paragraph("Enabled plugins: " + ", ".join(str(item) for item in selected_plugins))
            if selected_filters:
                document.add_paragraph("Enabled filters: " + ", ".join(str(item) for item in selected_filters))
            if attached_modules:
                _add_table(
                    document,
                    ["Module", "Kind", "Framework", "Power", "Capabilities"],
                    [
                        [
                            str(row.get("id", "module")),
                            str(row.get("kind", "-")),
                            str(row.get("framework", "-")),
                            str(row.get("power_score", 0)),
                            ", ".join(str(item) for item in list(row.get("capabilities", []) or [])[:4]),
                        ]
                        for row in attached_modules[:18]
                    ],
                )

        ocr_tooling = payload.get("ocr_tooling", {}) if isinstance(payload.get("ocr_tooling"), dict) else {}
        if ocr_tooling:
            pytesseract_info = ocr_tooling.get("pytesseract", {}) if isinstance(ocr_tooling.get("pytesseract"), dict) else {}
            document.add_heading("OCR Tooling", level=1)
            _add_table(
                document,
                ["Tool", "Available", "Details"],
                [
                    ["Preferred engine", "yes", str(ocr_tooling.get("preferred_engine", "none"))],
                    ["Pillow", str(bool((ocr_tooling.get("pillow") or {}).get("available"))).lower(), "image preprocessing/runtime"],
                    ["EasyOCR", str(bool((ocr_tooling.get("easyocr") or {}).get("available"))).lower(), "neural OCR backend"],
                    [
                        "PyTesseract",
                        str(bool(pytesseract_info.get("available"))).lower(),
                        f"binary_found={bool(pytesseract_info.get('tesseract_binary_found'))} path={pytesseract_info.get('tesseract_binary', '-') or '-'}",
                    ],
                ],
            )

        results = _safe_rows(payload.get("results"))
        document.add_heading("Results", level=1)
        if results:
            _add_table(
                document,
                ["Platform", "Status", "Confidence", "Response ms", "URL"],
                [
                    [
                        str(row.get("platform", "platform")),
                        str(row.get("status", "-")),
                        str(row.get("confidence", "-")),
                        str(row.get("response_time_ms", "-")),
                        str(row.get("url", "-")),
                    ]
                    for row in results[:28]
                ],
            )
        else:
            document.add_paragraph("No direct result rows were recorded.")

        issues = _safe_rows(payload.get("issues"))
        document.add_heading("Vulnerabilities", level=1)
        if issues:
            _add_table(
                document,
                ["Severity", "Title", "Scope", "Evidence", "Recommendation"],
                [
                    [
                        str(row.get("severity", "INFO")),
                        str(row.get("title", "Issue")),
                        str(row.get("scope", "-")),
                        str(row.get("evidence", "-"))[:120],
                        str(row.get("recommendation", "-"))[:120],
                    ]
                    for row in issues[:24]
                ],
            )
        else:
            document.add_paragraph("No exposure findings were reported.")

        plugins = _safe_rows(payload.get("plugins"))
        if plugins:
            document.add_heading("Plugin Intelligence", level=1)
            _add_table(
                document,
                ["Plugin", "Severity", "Summary", "Signals"],
                [
                    [
                        str(row.get("title", row.get("id", "plugin"))),
                        str(row.get("severity", "INFO")),
                        str(row.get("summary", "-"))[:140],
                        ", ".join(str(item) for item in list(row.get("highlights", []) or [])[:3]) or "-",
                    ]
                    for row in plugins[:18]
                ],
            )

        filters = _safe_rows(payload.get("filters"))
        if filters:
            document.add_heading("Filter Intelligence", level=1)
            _add_table(
                document,
                ["Filter", "Severity", "Summary", "Signals"],
                [
                    [
                        str(row.get("title", row.get("id", "filter"))),
                        str(row.get("severity", "INFO")),
                        str(row.get("summary", "-"))[:140],
                        ", ".join(str(item) for item in list(row.get("highlights", []) or [])[:3]) or "-",
                    ]
                    for row in filters[:18]
                ],
            )

        ocr_scan = payload.get("ocr_scan", {}) if isinstance(payload.get("ocr_scan"), dict) else {}
        ocr_items = _safe_rows(ocr_scan.get("items"))
        if ocr_items:
            document.add_heading("OCR Scan Highlights", level=1)
            _add_table(
                document,
                ["Source", "Kind", "Engine", "Confidence", "Signals"],
                [
                    [
                        str(row.get("source", "image")),
                        str(row.get("source_kind", "-")),
                        str(row.get("ocr_engine", "none")),
                        str(row.get("confidence_hint", "-")),
                        ", ".join(
                            f"{key}:{len(value)}"
                            for key, value in (row.get("signals", {}) or {}).items()
                            if isinstance(value, list) and value
                        )
                        or "-",
                    ]
                    for row in ocr_items[:20]
                ],
            )

        intelligence_bundle = payload.get("intelligence_bundle", {}) if isinstance(payload.get("intelligence_bundle"), dict) else {}
        if intelligence_bundle:
            metadata = intelligence_bundle.get("metadata", {}) if isinstance(intelligence_bundle.get("metadata"), dict) else {}
            confidence_distribution = (
                intelligence_bundle.get("confidence_distribution", {})
                if isinstance(intelligence_bundle.get("confidence_distribution"), dict)
                else {}
            )
            document.add_heading("Intelligence Summary", level=1)
            _add_table(
                document,
                ["Metric", "Value"],
                [
                    ["Entity count", str(metadata.get("entity_count", 0))],
                    ["Evidence count", str(metadata.get("evidence_count", 0))],
                    ["High confidence", str(confidence_distribution.get("high", 0))],
                    ["Medium confidence", str(confidence_distribution.get("medium", 0))],
                    ["Low confidence", str(confidence_distribution.get("low", 0))],
                ],
            )
    finally:
        if chart_dir is not None:
            chart_dir.cleanup()

    document.save(str(path))
    return str(path)
